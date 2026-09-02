#!/usr/bin/env python3
"""Render a migration markdown report as a formatted .docx.

Handles the subset of markdown used in docs/migration/*.md: ATX headings,
pipe tables, bullet and numbered lists, blockquote-free paragraphs, inline
**bold**, *italic*, `code` and [text](url) links, and horizontal rules.

Usage: python3 scripts/md-to-docx.py <input.md> [output.docx]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|\*[^*]+\*)")


def add_runs(par, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            par.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif part.startswith("[") and "](" in part:
            label, url = part[1:-1].split("](", 1)
            r = par.add_run(f"{label} ({url})")
            r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            r.underline = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            par.add_run(part[1:-1]).italic = True
        else:
            par.add_run(part)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator(line):
    return bool(re.fullmatch(r"\|[\s:\-|]+\|", line.strip()))


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell in enumerate(header):
        p = t.cell(0, j).paragraphs[0]
        add_runs(p, cell)
        for r in p.runs:
            r.bold = True
    for i, row in enumerate(body, start=1):
        for j in range(len(header)):
            p = t.cell(i, j).paragraphs[0]
            add_runs(p, row[j] if j < len(row) else "")
            if j and re.fullmatch(r"[\s$%\d.,*—–-]*", row[j] if j < len(row) else ""):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()


BLOCK_START = re.compile(r"(#{1,6}\s|[-*]\s|\d+\.\s|\||-{3,}$)")


def unwrap(lines):
    """Join soft-wrapped continuation lines into their preceding block."""
    out: list[str] = []
    for raw in lines:
        s = raw.strip()
        if (out and s and out[-1].strip() and not BLOCK_START.match(s)
                and not out[-1].strip().startswith("|")
                and not re.fullmatch(r"-{3,}", out[-1].strip())
                and not out[-1].strip().startswith("#")):
            out[-1] = out[-1].rstrip() + " " + s
        else:
            out.append(raw)
    return out


def convert(src: Path, dst: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    lines = unwrap(src.read_text().splitlines())
    i = 0
    table: list[list[str]] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("|"):
            if not is_separator(stripped):
                table.append(split_row(stripped))
            i += 1
            continue
        if table:
            add_table(doc, table)
            table = []

        if not stripped:
            i += 1
            continue
        if re.fullmatch(r"-{3,}", stripped):
            doc.add_paragraph().add_run().add_break()
            i += 1
            continue
        m = re.match(r"(#{1,6})\s+(.*)", stripped)
        if m:
            doc.add_heading(m.group(2), level=min(len(m.group(1)), 4))
            i += 1
            continue
        m = re.match(r"[-*]\s+(.*)", stripped)
        if m:
            add_runs(doc.add_paragraph(style="List Bullet"), m.group(1))
            i += 1
            continue
        m = re.match(r"\d+\.\s+(.*)", stripped)
        if m:
            add_runs(doc.add_paragraph(style="List Number"), m.group(1))
            i += 1
            continue
        # continuation lines of the previous list item / paragraph
        add_runs(doc.add_paragraph(), stripped)
        i += 1
    if table:
        add_table(doc, table)

    doc.save(dst)
    print(f"wrote {dst}")


if __name__ == "__main__":
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix(".docx")
    convert(src, dst)
